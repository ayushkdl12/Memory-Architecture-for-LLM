"""Generate docs/user_profile.docx — a rich user profile used to seed AI memory.

Run from backend/:  ./.venv/bin/python scripts/make_profile_docx.py
Output: ../docs/user_profile.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent.parent / "docs" / "user_profile.docx"

SECTIONS: dict[str, list[str]] = {
    "Personal": [
        "My name is Bibek.",
        "I am 26 years old.",
        "I live in Kathmandu, Nepal.",
        "My email address is bibek.kandel@example.com.",
        "My phone number is +977-9800000000.",
        "I prefer to be called Bibek.",
    ],
    "Work": [
        "I work as a software engineer.",
        "I work at TechNest.",
        "I have been a software engineer for 4 years.",
        "My team works on backend services.",
        "My manager is Priya.",
    ],
    "Technical skills": [
        "I am fluent in Python.",
        "I am fluent in Go.",
        "I know TypeScript.",
        "I know React.",
        "I work with PostgreSQL every day.",
        "I use Docker for local development.",
        "I am learning Kubernetes.",
    ],
    "Preferences": [
        "I prefer dark mode in all my tools.",
        "I prefer coding with a keyboard over a mouse.",
        "I am a morning person.",
        "I prefer black coffee.",
        "I prefer asynchronous communication.",
        "I like to listen to electronic music while working.",
    ],
    "Goals": [
        "My goal is to become a senior backend engineer.",
        "My goal is to publish a technical blog post this quarter.",
        "My goal is to complete the Kubernetes certification.",
        "My goal is to run a half marathon this year.",
    ],
    "Projects and deadlines": [
        "The project deadline for the memory agent project is July 10.",
        "The project deadline for the payments API is August 15.",
        "My performance review is scheduled for September 1.",
    ],
    "Upcoming events": [
        "I will attend the PyCon Nepal conference on October 12.",
        "I have a team standup every weekday at 10am.",
        "My birthday is on March 3.",
    ],
}


def build() -> Document:
    doc = Document()

    title = doc.add_heading("User Profile — Bibek", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "This document contains the durable profile of the user Bibek. It is "
        "used to seed and enrich the assistant's memory. Each bullet below "
        "states exactly one fact."
    )

    for heading, facts in SECTIONS.items():
        doc.add_heading(heading, level=1)
        for f in facts:
            p = doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("About me", level=1)
    about = doc.add_paragraph(
        "I am a backend software engineer from Kathmandu who enjoys building "
        "systems with Python and Go. I care about clean architecture, clear "
        "communication, and shipping on time. I drink too much coffee and I "
        "still prefer dark mode everywhere."
    )

    return doc


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
